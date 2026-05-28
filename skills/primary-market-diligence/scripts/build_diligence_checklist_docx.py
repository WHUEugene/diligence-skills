#!/usr/bin/env python3
"""Build the approved 商务尽调资料准备清单 DOCX from a structured JSON payload.

Input JSON schema:
{
  "metadata": {
    "委托单位": "...",
    "受托尽调机构": "...",
    "尽调对象": "...",
    "适用阶段": "..."
  },
  "sections": [
    {
      "title": "第一类：企业主体与背景信息",
      "rows": [
        {
          "序号": "1.1",
          "需了解的核心问题": "...",
          "推荐提供的资料": "正式资料：...；替代资料：...。",
          "核心目的": "...",
          "资料提供状态": ""
        }
      ]
    }
  ]
}
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


STANDARD_SECTIONS = [
    "第一类：企业主体与背景信息",
    "第二类：项目技术与产品基础",
    "第三类：市场规模与竞争",
    "第四类：项目投资与建设方案",
    "第五类：资金筹措与风险保障",
    "第六类：过往业绩与商誉",
    "第七类：拟派管理团队与组织规划",
    "第八类：财务预测与经营假设",
    "第九类：工商登记与关联方信息",
    "第十类：诉讼、行政处罚与合规记录",
    "第十一类：上市进展与重大资本运作",
    "第十二类：行业政策与竞争动态",
]

HEADERS = ["序号", "需了解的核心问题", "推荐提供的资料", "核心目的", "资料提供状态"]
META_KEYS = ["委托单位", "受托尽调机构", "尽调对象", "适用阶段"]


def set_cell_shading(cell, fill: str = "FFFFFF") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = "BFBFBF", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = mar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_width(table, col_widths_cm: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    tbl.insert(1, grid)

    widths_dxa = [int(Cm(w).emu / 635) for w in col_widths_cm]
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    for row in table.rows:
        for idx, width in enumerate(col_widths_cm):
            cell = row.cells[idx]
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))


def set_run_font(run, size: float, bold: bool = False, color: str | None = None, font: str = "楷体") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size: float, bold: bool = False, align=None, font: str = "楷体") -> None:
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, font=font)


def set_cell_text(cell, text: str, size: float = 9.4, bold: bool = False, align=None) -> None:
    cell.text = ""
    parts = str(text or "").split("\n")
    for idx, part in enumerate(parts):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        run = p.add_run(part)
        set_run_font(run, size=size, bold=bold)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.08
        if align is not None:
            p.alignment = align


def add_metadata(doc: Document, metadata: dict) -> None:
    for key in META_KEYS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.25
        label = p.add_run(f"{key}：")
        set_run_font(label, size=11, bold=True)
        value = p.add_run(str(metadata.get(key, "") or ""))
        set_run_font(value, size=11)


def add_section_table(doc: Document, title: str, rows: list[dict]) -> None:
    h = doc.add_paragraph()
    r = h.add_run(title)
    set_run_font(r, size=12, bold=True, color="000000")
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(8)

    table = doc.add_table(rows=1, cols=len(HEADERS))
    # Match the approved wide-table layout: narrow serial/status, broad middle columns.
    set_table_width(table, [1.20, 6.75, 10.65, 5.25, 1.75])

    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    set_row_cant_split(header_row)
    for i, header in enumerate(HEADERS):
        cell = header_row.cells[i]
        set_cell_shading(cell, "FFFFFF")
        set_cell_borders(cell, color="9E9E9E", size="4")
        set_cell_margins(cell, top=120, bottom=120, start=120, end=120)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_text(cell, header, size=10.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for item in rows:
        row = table.add_row()
        set_row_cant_split(row)
        values = [
            item.get("序号", ""),
            item.get("需了解的核心问题", ""),
            item.get("推荐提供的资料", ""),
            item.get("核心目的", ""),
            item.get("资料提供状态", "") or "",
        ]
        for i, value in enumerate(values):
            cell = row.cells[i]
            set_cell_borders(cell)
            set_cell_margins(cell, top=95, bottom=95, start=130, end=130)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            align = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 4) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cell, value, size=9.4, align=align)


def validate_payload(payload: dict) -> None:
    sections = payload.get("sections") or []
    titles = [s.get("title", "") for s in sections]
    missing = [s for s in STANDARD_SECTIONS if s not in titles]
    if missing:
        raise ValueError(f"Missing standard sections: {missing}")
    if len(sections) != 12:
        raise ValueError(f"Expected 12 sections, got {len(sections)}")
    for section in sections:
        rows = section.get("rows") or []
        if not rows:
            raise ValueError(f"Section has no rows: {section.get('title')}")
        for row in rows:
            for header in HEADERS[:4]:
                if not str(row.get(header, "")).strip():
                    raise ValueError(f"Missing {header} in section {section.get('title')}")


def build_docx(payload: dict, output: Path) -> None:
    validate_payload(payload)

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.left_margin = Cm(1.9)
    sec.right_margin = Cm(1.9)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "楷体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "楷体")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "楷体")
    normal.font.size = Pt(9.4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(22)
    run = title.add_run("商务尽调资料准备清单")
    set_run_font(run, size=17, bold=True)

    add_metadata(doc, payload.get("metadata") or {})

    for section_title in STANDARD_SECTIONS:
        section = next(s for s in payload["sections"] if s.get("title") == section_title)
        add_section_table(doc, section_title, section.get("rows") or [])

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, help="Input JSON payload path")
    parser.add_argument("--output", required=True, help="Output DOCX path")
    args = parser.parse_args()

    payload = json.loads(Path(args.input).read_text(encoding="utf-8"))
    build_docx(payload, Path(args.output))
    print(args.output)


if __name__ == "__main__":
    main()
