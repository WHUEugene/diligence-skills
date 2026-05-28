#!/usr/bin/env python3
"""Extract a reusable formal-report master template from a reference DOCX."""

from __future__ import annotations

import argparse
import json
import re
import shutil
from collections import Counter
from pathlib import Path

from docx import Document


FRONT_MATTER_TERMS = ("重要声明", "重大事项提示", "目录", "释义")
SOURCE_TERMS = ("资料来源", "数据来源", "来源")


def paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def classify_paragraph(text: str, style: str) -> str:
    compact = text.replace(" ", "")
    if style.startswith("Heading"):
        return "heading"
    if any(term in compact for term in FRONT_MATTER_TERMS):
        return "front_matter"
    if re.match(r"^(图|表)\s*\d", compact):
        return "caption"
    if any(term in compact for term in SOURCE_TERMS):
        return "source_note"
    if re.search(r"\d+(?:\.\d+)?\s*(?:亿元|万元|%|GWh|μm|um|平方米|吨|年|个月)", text):
        return "data_paragraph"
    return "body"


def collect_heading_records(doc: Document) -> list[dict]:
    records: list[dict] = []
    for index, paragraph in enumerate(doc.paragraphs, 1):
        text = normalize(paragraph_text(paragraph))
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        if style.startswith("Heading") or any(term in text.replace(" ", "") for term in FRONT_MATTER_TERMS):
            records.append(
                {
                    "paragraph_index": index,
                    "style": style,
                    "title": text,
                    "role": classify_paragraph(text, style),
                }
            )
    return records


def collect_table_records(doc: Document) -> list[dict]:
    records: list[dict] = []
    for idx, table in enumerate(doc.tables, 1):
        rows = len(table.rows)
        cols = len(table.columns)
        first_row = []
        if rows:
            first_row = [normalize(cell.text) for cell in table.rows[0].cells]
        records.append(
            {
                "table_no": idx,
                "rows": rows,
                "cols": cols,
                "header": first_row,
                "slot_type": infer_table_slot(first_row),
            }
        )
    return records


def infer_table_slot(header: list[str]) -> str:
    joined = " ".join(header)
    if any(term in joined for term in ("产品", "应用", "指标")):
        return "product_or_technology"
    if any(term in joined for term in ("收入", "毛利", "投资", "金额", "财务")):
        return "finance_or_investment"
    if any(term in joined for term in ("公司", "可比", "竞争", "客户")):
        return "peer_or_market"
    if any(term in joined for term in ("风险", "条件", "核验", "资料")):
        return "risk_or_evidence"
    return "general"


def collect_style_summary(doc: Document) -> dict:
    para_styles = Counter((p.style.name if p.style else "") for p in doc.paragraphs)
    table_count = len(doc.tables)
    sections = []
    for section in doc.sections:
        sections.append(
            {
                "page_width": section.page_width,
                "page_height": section.page_height,
                "top_margin": section.top_margin,
                "bottom_margin": section.bottom_margin,
                "left_margin": section.left_margin,
                "right_margin": section.right_margin,
            }
        )
    return {
        "paragraph_styles": dict(para_styles.most_common()),
        "table_count": table_count,
        "sections": sections,
    }


def build_slot_map(headings: list[dict], tables: list[dict]) -> list[dict]:
    slots: list[dict] = []
    current_chapter = ""
    table_iter = iter(tables)
    pending_tables = list(table_iter)
    table_idx = 0
    for heading in headings:
        title = heading["title"]
        style = heading["style"]
        if style == "Heading 1" or title in ("重大事项提示", "释义"):
            current_chapter = title
        evidence_need = infer_evidence_need(title)
        slot = {
            "chapter": current_chapter or title,
            "title": title,
            "style": style,
            "role": heading["role"],
            "evidence_need": evidence_need,
            "reuse_rule": "same_structure_fill_with_new_project_evidence",
        }
        if table_idx < len(pending_tables) and (
            "表" in title or any(term in title for term in ("产品", "财务", "投资", "风险", "资料"))
        ):
            slot["nearby_table_hint"] = pending_tables[table_idx]
            table_idx += 1
        slots.append(slot)
    return slots


def infer_evidence_need(title: str) -> str:
    compact = title.replace(" ", "")
    if any(term in compact for term in ("重要声明", "目录", "释义")):
        return "template_text_or_definition"
    if any(term in compact for term in ("行业", "市场", "竞争格局", "产业链")):
        return "public_sources_and_comparable_filings"
    if any(term in compact for term in ("主体", "股权", "历史沿革", "资质", "知识产权")):
        return "target_legal_and_registry_evidence"
    if any(term in compact for term in ("业务", "技术", "产品", "研发", "生产", "客户", "采购")):
        return "target_materials_plus_external_validation"
    if any(term in compact for term in ("财务", "预测", "收入", "毛利", "投资规模")):
        return "financial_model_and_supporting_schedules"
    if any(term in compact for term in ("风险", "结论", "建议", "风控")):
        return "evidence_synthesis_and_transaction_terms"
    return "project_or_public_evidence"


def write_markdown(path: Path, data: dict) -> None:
    lines = [
        "# 正式报告母版模板",
        "",
        f"- 参考文件：{data['reference_docx']}",
        f"- 段落数：{data['paragraph_count']}",
        f"- 表格数：{len(data['tables'])}",
        f"- 标题/前置节点数：{len(data['headings'])}",
        "",
        "## 母版使用规则",
        "",
        "1. 有新项目材料时，复制本母版的前置结构、目录节奏、章节层级、标题样式、表格样式和图表节奏。",
        "2. 只把新项目材料、公开资料和可比招股书能支持的事实填入对应槽位。",
        "3. 客户、订单、股权、设备、财务预测、检测报告等私有证据缺失时，不写成确定结论；转为风险、先决条件或待核验事项。",
        "4. 报告正文不得出现“参考母版”“复制正式报告”“由于只有PPT”等过程语言。",
        "",
        "## 章节槽位",
    ]
    for slot in data["slots"]:
        lines.append(
            f"- `{slot['style'] or 'Normal'}` {slot['title']} | 证据要求：{slot['evidence_need']} | 复用：{slot['reuse_rule']}"
        )
    lines += ["", "## 表格槽位"]
    for table in data["tables"]:
        header = " / ".join(table["header"])
        lines.append(f"- 表{table['table_no']}: {table['rows']}行 x {table['cols']}列 | {table['slot_type']} | {header}")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("reference_docx")
    parser.add_argument("--output-dir", required=True)
    args = parser.parse_args()

    reference = Path(args.reference_docx).expanduser().resolve()
    out = Path(args.output_dir).expanduser().resolve()
    out.mkdir(parents=True, exist_ok=True)

    doc = Document(reference)
    headings = collect_heading_records(doc)
    tables = collect_table_records(doc)
    data = {
        "reference_docx": str(reference),
        "paragraph_count": len(doc.paragraphs),
        "headings": headings,
        "tables": tables,
        "style_summary": collect_style_summary(doc),
        "slots": build_slot_map(headings, tables),
    }

    (out / "formal_report_master_template.json").write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    write_markdown(out / "formal_report_master_template.md", data)
    shutil.copy2(reference, out / "formal_report_style_master.docx")
    print(json.dumps({"ok": True, "output_dir": str(out), "headings": len(headings), "tables": len(tables)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
