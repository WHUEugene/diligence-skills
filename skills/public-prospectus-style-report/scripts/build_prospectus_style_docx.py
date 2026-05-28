#!/usr/bin/env python3
"""Build a prospectus-style DOCX report from a Markdown intermediate draft.

This generator intentionally supports a compact subset of Markdown used by the
public-prospectus-style-report skill: headings, paragraphs, ordered/unordered
lists, and simple pipe tables. It rejects meta/preface sections that should not
appear in the formal deliverable.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


FORBIDDEN_HEADING_PATTERNS = [
    "重要提示",
    "重大事项提示",
    "重要声明",
    "报告边界",
    "使用说明",
    "附录",
]

EXPECTED_H2_HEADINGS = [
    "一、行业情况",
    "二、项目业务情况",
    "三、项目在行业环境中的优劣势",
    "四、投资方案与风控措施",
    "资料来源与待核验事项",
]

CITATION_RE = re.compile(r"(\[[A-Z]{1,5}\d+\])")
IMAGE_RE = re.compile(r"^!\[(.*)\]\(([^)]+)\)\s*$")


def set_run_font(run, size: float, bold: bool = False, font: str = "宋体", color: str = "000000") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_format(paragraph, first_line: bool = False) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.28
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)


def set_cell_borders(cell, color: str = "BFBFBF", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str = "FFFFFF") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def clean_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("__", "").replace("`", "")
    return text.strip()


def add_inline_runs(paragraph, text: str, size: float = 10.5, bold: bool = False, font: str = "宋体") -> None:
    """Add body text and render source markers such as [P1] as small superscript."""

    for part in CITATION_RE.split(clean_inline(text)):
        if not part:
            continue
        is_citation = bool(CITATION_RE.fullmatch(part))
        run = paragraph.add_run(part)
        set_run_font(run, size=7.5 if is_citation else size, bold=bold and not is_citation, font=font)
        if is_citation:
            run.font.superscript = True


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    add_inline_runs(p, text.strip(), size=9.2, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_borders(cell)
    set_cell_shading(cell, "FFFFFF")


def is_table_start(lines: list[str], idx: int) -> bool:
    if idx + 1 >= len(lines):
        return False
    return lines[idx].strip().startswith("|") and re.match(r"^\s*\|?[-:\s|]+\|?\s*$", lines[idx + 1])


def parse_table(lines: list[str], idx: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        line = lines[idx].strip().strip("|")
        if re.match(r"^[-:\s|]+$", line):
            idx += 1
            continue
        rows.append([cell.strip() for cell in line.split("|")])
        idx += 1
    return rows, idx


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            text = row[c_idx] if c_idx < len(row) else ""
            set_cell_text(table.cell(r_idx, c_idx), text, bold=(r_idx == 0))
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int) -> None:
    if any(pattern in text for pattern in FORBIDDEN_HEADING_PATTERNS):
        raise ValueError(f"Forbidden formal-report heading: {text}")
    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(18)
        add_inline_runs(p, text, size=18, bold=True, font="黑体")
        return
    if level == 2:
        if doc.paragraphs:
            doc.add_section(WD_SECTION.NEW_PAGE)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        add_inline_runs(p, text, size=14, bold=True, font="黑体")
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    add_inline_runs(p, text, size=11.5, bold=True, font="黑体")


def add_image(doc: Document, base_dir: Path, alt_text: str, image_ref: str) -> None:
    image_path = Path(image_ref)
    if not image_path.is_absolute():
        image_path = base_dir / image_path
    if not image_path.exists():
        p = doc.add_paragraph()
        set_paragraph_format(p)
        add_inline_runs(p, f"[图表文件缺失：{image_ref}]", size=10.5)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(13.2))
    if alt_text:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(8)
        caption_text = alt_text if alt_text.startswith("图") else f"图：{alt_text}"
        add_inline_runs(caption, caption_text, size=9.2, font="宋体")


def audit_four_module_headings(markdown: str) -> None:
    h2 = []
    for line_no, line in enumerate(markdown.splitlines(), 1):
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if not match:
            continue
        level = len(match.group(1))
        title = match.group(2).strip()
        if any(pattern in title for pattern in FORBIDDEN_HEADING_PATTERNS):
            raise ValueError(f"Forbidden formal-report heading at line {line_no}: {title}")
        if level == 2:
            h2.append(title)
    if h2 != EXPECTED_H2_HEADINGS:
        raise ValueError(
            "Four-module report headings must be exactly: "
            + " / ".join(EXPECTED_H2_HEADINGS)
            + f"; got: {' / '.join(h2)}"
        )


def build_docx(markdown: str, output: Path, require_four_module: bool = True, base_dir: Path | None = None) -> None:
    if require_four_module:
        audit_four_module_headings(markdown)
    base_dir = base_dir or Path.cwd()
    lines = markdown.splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)

    idx = 0
    while idx < len(lines):
        raw = lines[idx]
        line = raw.strip()
        if not line:
            idx += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            add_heading(doc, heading.group(2), len(heading.group(1)))
            idx += 1
            continue

        image = IMAGE_RE.match(line)
        if image:
            add_image(doc, base_dir, image.group(1), image.group(2))
            idx += 1
            continue

        if is_table_start(lines, idx):
            table_rows, idx = parse_table(lines, idx)
            add_table(doc, table_rows)
            continue

        list_match = re.match(r"^(\d+\.\s+|[-*]\s+)(.+)$", line)
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line=not bool(list_match))
        prefix = ""
        content = line
        if list_match:
            marker = list_match.group(1)
            prefix = "• " if marker.strip() in {"-", "*"} else marker
            content = list_match.group(2)
        add_inline_runs(p, prefix + content, size=10.5)
        idx += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, help="Markdown draft path")
    parser.add_argument("--output", required=True, help="Output DOCX path")
    parser.add_argument(
        "--allow-non-four-module",
        action="store_true",
        help="Skip the leader four-module structure gate. Use only for explicit full-prospectus simulations.",
    )
    args = parser.parse_args()
    input_path = Path(args.input)
    build_docx(
        input_path.read_text(encoding="utf-8"),
        Path(args.output),
        require_four_module=not args.allow_non_four_module,
        base_dir=input_path.parent,
    )
    print(args.output)


if __name__ == "__main__":
    main()
